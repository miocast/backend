import array
from itertools import count
import os
from pathlib import Path
import re
import requests
import urllib3
import json
from bs4 import BeautifulSoup

from docx import Document

from app.business.document import DocumentModel
from app.config.settings import MAIN_API
from app.services.llm_service import LlmService

class ActParser():
    def __init__(self):
        urllib3.disable_warnings()
        self.llm_service = LlmService()

    def parse_by_theme(self, theme: str):
        request_str  = "{\"query\":\"{" + theme + "}\",\"type\":\"EXACT\",\"mode\":\"SIMPLE\"}"

        json_data = {
            "request": {
                "filter": None,
                "groups": [
                    "Текущие редакции"
                ],
                "dateFrom": None,
                "dateTo": None,
                "sortOrder": "desc",
                "sortField": "document_date_edition",
                "groupField": None,
                "type": "MULTIQUERY",
                "multiqueryRequest": {
                    "queryRequests": [
                        {
                            "type": "Q",
                            "queryRequestRole": "SIMPLE",
                            "request": request_str,
                            "operator": "AND"
                        },
                        {
                            "type": "SQ",
                            "queryId": "15556f4a-7811-42a5-a976-a9eb9e90cb74",
                            "operator": "AND"
                        }
                    ]
                },
                "simpleSearchFieldsBundle": "test1",
                "start": 0,
                "rows": 10,
                "uid": "07f51374-78a0-4df6-a243-92c9e01f8c1a",
                "noOrpho":  False,
                "facet": {
                    "field": [
                        "type"
                    ]
                },
                "facetLimit": 21,
                "additionalFields": [
                    "document_name",
                    "document_normative",
                    "document_state"
                ],
                "groupLimit": 3,
                "woBoost": False
            },
            "doNotSaveHistory": False
        }

        try:
        #todo: Тут парсить сайтик. Сохранять в куда то. Отправлять на шарпы наименованаия
            response = requests.post('https://pravo-search.minjust.ru/bigs/s.action', json=json_data, verify=False)

            data = response.json()
            data = data["searchResult"]

            if (len(data["documents"]) == 0):
                return
            
            documents_info: list[DocumentModel] = []
            
            for doc in data["documents"]:
                id = doc["id"]
                name = str(doc["name"])
                name = re.sub(r'[<>:"/\\|?*]', '', name)
                type = doc["type"]

                folder = Path(__file__).parent.absolute()
                file_path = str(Path(__file__).parent.absolute().joinpath(f"{type}/{name}.docx"))

                os.makedirs(file_path, exist_ok=True)

                link = f"https://pravo-search.minjust.ru/bigs/showDocumentWithTemplate.action?id={id}&shard=Текущие%20редакции&templateName=printText.flt"
                doc_html = requests.get(link, verify=False)
                doc_text = doc_html.text
                
                soup = BeautifulSoup(doc_text, "html.parser")

                # documents_info.append({"id": id, "name": name, "type": type, "file_path": file_path, "doc_text": soup.text})
                documents_info.append(DocumentModel(id=id, name=name, type=type, file_path=file_path, text=re.sub(r'(^\s+|\n\s*\n)', '\n', soup.get_text(), flags=re.MULTILINE)))
                
                doc = Document()

        # Добавляем текст в документ
                for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol']):
                    if element.name in ['h1', 'h2', 'h3']:
                        # Заголовки
                        level = int(element.name[1])
                        doc.add_heading(element.get_text(), level=level)
                    elif element.name == 'ul':
                    # Маркированный список
                        for li in element.find_all('li'):
                            doc.add_paragraph(li.get_text(), style='List Bullet')
                    elif element.name == 'ol':
                            # Нумерованный список
                        for li in element.find_all('li'):
                            doc.add_paragraph(li.get_text(), style='List Number')
                    else:
                            # Обычный текст
                        doc.add_paragraph(element.get_text())

                    
                    # Сохраняем документ
                    doc.save(file_path)
        
            self.llm_service.save_docs_to_vector_db(documents_info)
            self._send_new_documents_to_api(documents_info)
        except Exception as ex:
            print(ex)

        # TODO: Отправить наименования в шарпы. Проставить статус готово на стороне шарпов
        
    def _send_new_documents_to_api(self, documents: list[DocumentModel]):
        dict = []
        
        for doc in documents:
            dict.append({
            "id": doc.id,
            "name": doc.name,
            "type": doc.type,
            "filePath": doc.file_path,
            })
        
        url = MAIN_API + "/api/v1/Npa/store-from-worker"    
        json_data = json.dumps(dict)
        print(json_data)
        try:
            headers = {
    "Content-Type": "application/json",}
            response = requests.post(url, data=json.dumps(dict), headers=headers)
            print(response.json())
        except Exception as e:
            print(e)